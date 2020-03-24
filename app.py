from flask import Flask
from flask import redirect
from flask import render_template
from flask import request
from flask import session
from flask import url_for
from flask import flash
from flask import send_from_directory
from flask_bcrypt import Bcrypt
from flask_bcrypt import generate_password_hash
from flask_wtf import FlaskForm, CSRFProtect
from wtforms import TextField, PasswordField, TextAreaField, validators, SelectField, StringField, SubmitField
from wtforms.validators import DataRequired
from wtforms.widgets import TextArea
from werkzeug.utils import secure_filename
from flask_migrate import Migrate
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy.orm.exc import NoResultFound
from pytz import timezone
from docx import Document

import uuid
import pytz
import os
import shutil
from pathlib import Path
import logging
import calendar
import config
import random
import datetime
import time
import re
import csv
import json
import requests
from collections import OrderedDict
#Debugging Tools
import inspect

# logger stuff
logger = logging.getLogger(__name__)
formatter = logging.Formatter(
    '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
console = logging.StreamHandler()
console.setLevel(logging.DEBUG)
console.setFormatter(formatter)
logger.addHandler(console)

# init app and load conf
app = Flask(__name__, static_url_path='/templates')
app.config.from_object(config)
csrf = CSRFProtect(app)
bcrypt = Bcrypt(app)

# init db
db = SQLAlchemy(app)
migrate = Migrate(app, db)

# init global
tteconvention_data = {}


# -----------------------------------------------------------------------
#  Classes
# -----------------------------------------------------------------------
class Convention:
    def __init__(self, name):
        self.name = name
    def add_location(self, geolocation_name):
        self.location = geolocation_name
    def add_description(self, description):
        self.description = description
    def add_phone_number(self, phone_number):
        self.phone_number = phone_number
    def add_email(self, email):
        self.email = email
    def add_dates(self, dates):
        self.dates = dates

# -----------------------------------------------------------------------
# Database models
# -----------------------------------------------------------------------
class Conventions(db.Model):
    name = db.Column(db.String(255))
    tteid = db.Column(db.String(255), primary_key=True)
    slots = db.Column(db.String(2048))
    tables = db.Column(db.String(255))

class Volunteers(db.Model):
    name = db.Column(db.String(255))
    email = db.Column(db.String(255))
    role = db.Column(db.String(255))
    hours = db.Column(db.Integer)
    tiers = db.Column(db.String(255))
    slots = db.Column(db.String(255))
    conventions = db.Column(db.String(255))
    tteid = db.Column(db.String(255), primary_key=True)

class User(db.Model):
    id = db.Column(db.String(255), primary_key=True)
    name = db.Column(db.String(255))
    email = db.Column(db.String(255))
    role = db.Column(db.String(255))
    username = db.Column(db.String(255))
    password = db.Column(db.String(255))

# -----------------------------------------------------------------------
# Forms
# -----------------------------------------------------------------------
class LoginForm(FlaskForm):
    username = StringField('username:', validators=[validators.DataRequired()])
    password = PasswordField('Password:', validators=[validators.DataRequired()])

class FileForm(FlaskForm):
    selectfile = SelectField('Filename', validators=[validators.DataRequired()])
    volunteerreport = SubmitField(label='Create Report for Volunteers')
    volunteercsv = SubmitField(label='Create CSV file for Volunteers')
    eventcsv = SubmitField(label='Create CSV file for Volunteers')
    eventsave = SubmitField(label='Submit File for Convention Events')
    conventionsave = SubmitField(label='Submit File for Convention Details')
    eventsdelete = SubmitField(label='Delete All Convention Events')
    shiftsdelete = SubmitField(label='Delete All Volunteer Shifts ')
    daypartsdelete = SubmitField(label='Delete All Convention Day Parts')
    roomsandtablesdelete = SubmitField(label='Delete All Convention Rooms and Tables')
    volunteerdelete = SubmitField(label='Delete Volunteer')

class ConForm(FlaskForm):
    selectcon = SelectField('Convention', validators=[validators.DataRequired()])
    consubmit = SubmitField(label='Submit')

class LogoutForm(FlaskForm):
    logoutsubmit = SubmitField(label='Logout')

class NewConventionForm(FlaskForm):
    name = StringField('New Convention Name:', validators=[validators.DataRequired()])
    location = StringField('City, State of the Convention:', validators=[validators.DataRequired()])
    description = TextAreaField('Description of the Convention:', widget=TextArea(), validators=[validators.DataRequired()])
    phone_number = StringField('Please provide your phone number for volunteers to contact you: ', validators=[validators.DataRequired()])
    email = StringField('Please provide your email for volunteers to contact you: ', validators=[validators.DataRequired()])
    dates = TextAreaField('List each date of the Convention, one date per line:', widget=TextArea(), validators=[validators.DataRequired()])
    conventionsubmit = SubmitField(label='Submit')

class CreateUserForm(FlaskForm):
    name = StringField('Your Name:', validators=[validators.DataRequired()])
    username = StringField('username:', validators=[validators.DataRequired()])
    password = PasswordField('Password:', validators=[validators.DataRequired()])
    email = StringField('email address:', validators=[validators.DataRequired()])
    role = SelectField('User Role', choices=[('admin','admin'),('coodinator','coodinator')],validators=[validators.DataRequired()])

class ResetUserForm(FlaskForm):
    oldpassword = PasswordField('Old Password:', validators=[validators.DataRequired()])
    passwordcheck = PasswordField('Verify Old Password:', validators=[validators.DataRequired()])
    newpassword = PasswordField('New Password:', validators=[validators.DataRequired()])


# -----------------------------------------------------------------------
# Internal Functions
# -----------------------------------------------------------------------
# -----------------------------------------------------------------------
# Helper Functions
# -----------------------------------------------------------------------
# -----------------------------------------------------------------------
# Convert a given datetime object to a UTC time
# -----------------------------------------------------------------------
def datetime_utc_convert(ttesession,tteconvention_id,unconverted_datetime):
    timezone_data = tte_convention_geolocation_api_get(ttesession,tteconvention_id)
    current_tz = timezone(timezone_data)
    dt = current_tz.localize(unconverted_datetime)
    utc_delta = dt.utcoffset()
    utc_time = unconverted_datetime - utc_delta
    return(utc_time)

# -----------------------------------------------------------------------
# Convert UTC datetime object to the convention time
# -----------------------------------------------------------------------
def datetime_timezone_convert(ttesession,tteconvention_id, utc_datetime):
    timezone_data = tte_convention_geolocation_api_get(ttesession,tteconvention_id)
    current_tz = timezone(timezone_data)
    utc_tz = timezone('UTC')
    current_time = utc_tz.localize(utc_datetime).astimezone(current_tz)
    return(current_time)

# -----------------------------------------------------------------------
# Start a Session to TTE
# -----------------------------------------------------------------------
def tte_session():
    params = {'api_key_id': config.tte_api_key_id, 'username' : config.tte_username, 'password': config.tte_password}
    response = requests.post(config.tte_url + "/session", params=params)
    if response.status_code==200:
        session = response.json()['result']
    return (session)

# -----------------------------------------------------------------------
# Create the updateconform object to display pre-existing data for a convention
# -----------------------------------------------------------------------
def conform_info():
    all_days = []
    #Create and populate a new instance of the Convention class
    this_convention = Convention(tteconvention_data['result']['name'])
    this_convention.add_location(tteconvention_data['result']['geolocation_name'])
    this_convention.add_phone_number(tteconvention_data['result']['phone_number'])
    this_convention.add_email(tteconvention_data['result']['email_address'])
    this_convention.add_description(tteconvention_data['result']['description'])
    for day in tteconvention_data['result']['days']:
        dayonly = day['day_time'].strftime('%m/%d/%Y')
        all_days.append(dayonly)
    this_convention.add_dates(all_days)
    updateconform = NewConventionForm(request.form, obj=this_convention)
    updateconform.populate_obj(this_convention)
    return(updateconform)

# -----------------------------------------------------------------------
# Output a listing of events run by a specific host
# -----------------------------------------------------------------------
def create_volunteer_report(ttesession,tteconvention_id):
    document = Document()

    sorted_volunters = sorted(tteconvention_data['volunteers'],key = lambda j: j['lastname'])
    for volunteer in sorted_volunters:
        volunteer_events = []
        for event in tteconvention_data['events']:
            for host in event['hosts']:
                if volunteer['user_id'] == host['user_id']:
                    volunteer_events.append(event)
                    volunteer['events'] = volunteer_events
                else:
                    pass
        volunteer['shifts'] = tte_volunteer_shifts_api_get(ttesession,tteconvention_id,volunteer['id'])

        #Create volunteer page in doc with info from TTE
        document.add_heading(volunteer['name'], level=1)
        try:
            document.add_heading(volunteer['email_address'], level=2)
        except KeyError:
            pass
        document.add_heading('Volunteer Information',level=2)
        volunteer_table = document.add_table(rows=10, cols=2)

        volunteer_row_cells = volunteer_table.rows[0].cells
        volunteer_row_cells[0].text = 'Volunteer Location'
        try:
            volunteer_row_cells[1].text = volunteer['custom_fields']['volunteerlocation']
        except KeyError:
            pass
        volunteer_row_cells = volunteer_table.rows[1].cells
        volunteer_row_cells[0].text = 'Volunteer Pronouns'
        try:
            volunteer_row_cells[1].text = volunteer['custom_fields']['volunteerpronouns']
        except KeyError:
            pass

        volunteer_row_cells = volunteer_table.rows[2].cells
        volunteer_row_cells[0].text = 'Volunteer Level'
        try:
            volunteer_row_cells[1].text = volunteer['custom_fields']['volunteerlevel']
        except KeyError:
            pass

        volunteer_row_cells = volunteer_table.rows[3].cells
        volunteer_row_cells[0].text = 'Volunteer Source'
        try:
            volunteer_row_cells[1].text = volunteer['custom_fields']['volunteersource']
        except KeyError:
            pass

        volunteer_row_cells = volunteer_table.rows[4].cells
        volunteer_row_cells[0].text = 'Volunteer Emergency Contact'
        try:
            volunteer_row_cells[1].text = volunteer['custom_fields']['volunteeremergencycontact']
        except KeyError:
            pass

        volunteer_row_cells = volunteer_table.rows[5].cells
        volunteer_row_cells[0].text = 'Volunteer Experience'
        try:
            volunteer_row_cells[1].text = volunteer['custom_fields']['volunteerexperience']
        except KeyError:
            pass

        volunteer_row_cells = volunteer_table.rows[6].cells
        volunteer_row_cells[0].text = 'Volunteer Role(s)'
        try:
            volunteer_row_cells[1].text = volunteer['custom_fields']['volunteerrole']
        except KeyError:
            pass

        volunteer_row_cells = volunteer_table.rows[7].cells
        volunteer_row_cells[0].text = 'Volunteer Highest Tier'
        try:
            volunteer_row_cells[1].text = volunteer['custom_fields']['volunteertiers']
        except KeyError:
            pass

        volunteer_row_cells = volunteer_table.rows[8].cells
        volunteer_row_cells[0].text = 'Volunteer Shirt Size'
        try:
            volunteer_row_cells[1].text = volunteer['custom_fields']['volunteershirtsize']
        except KeyError:
            pass

        volunteer_row_cells = volunteer_table.rows[9].cells
        volunteer_row_cells[0].text = 'Volunteer Comments'
        try:
            volunteer_row_cells[1].text = volunteer['custom_fields']['volunteerother']
        except KeyError:
            pass

        # Create page for volunter with their shift information
        if len(volunteer['shifts']) != 0:
            document.add_page_break()
            document.add_heading(volunteer['name'], level=1)
            document.add_heading('Volunteer Shifts',level=2)
            shifts_table = document.add_table(rows=1, cols=3)
            shifts_hdr_cells = shifts_table.rows[0].cells
            shifts_hdr_cells[0].text = 'Shift Name'
            shifts_hdr_cells[1].text = 'Day'
            shifts_hdr_cells[2].text = 'Time Range'
            volunter_sorted_shifts = sorted(volunteer['shifts'],key = lambda i: i['shift_data']['start_time'])
            for vol_shift in volunter_sorted_shifts:
                shifts_row_cells = shifts_table.add_row().cells
                shifts_row_cells[0].text = vol_shift['shift_data']['name']
                shift_datetime_utc = datetime.datetime.strptime(vol_shift['shift_data']['start_time'], '%Y-%m-%d %H:%M:%S')
                shift_datetime_converted = datetime_timezone_convert(ttesession,tteconvention_id, shift_datetime_utc)
                vol_shift_day = shift_datetime_converted.strftime('%a %b %d')
                shifts_row_cells[1].text = vol_shift_day
                shifts_row_cells[2].text = vol_shift['shift_data']['times_range']

        # Create page for volunteer with info on the events they are hosting
        if len(volunteer_events) != 0:
            document.add_page_break()
            document.add_heading(volunteer['name'], level=1)
            document.add_heading('Event Information',level=2)
            events_table = document.add_table(rows=1, cols=5)
            events_hdr_cells = events_table.rows[0].cells
            events_hdr_cells[0].text = 'Event Name'
            events_hdr_cells[1].text = 'Duration'
            events_hdr_cells[2].text = 'Room'
            events_hdr_cells[3].text = 'Table'
            events_hdr_cells[4].text = 'Start Time'
            volunteer_sorted_events = sorted(volunteer_events, key = lambda k: k['start_date'])
            for vol_event in volunteer_sorted_events:
                events_row_cells = events_table.add_row().cells
                events_row_cells[0].text = vol_event['name']
                events_row_cells[1].text = str(vol_event['duration'])
                events_row_cells[2].text = vol_event['room_name']
                events_row_cells[3].text = vol_event['space_name']
                events_row_cells[4].text = vol_event['startdaypart_name']
        document.add_paragraph('Total Hours: ' + str(volunteer['hours_scheduled_count']) )
        document.add_page_break()

    doc_name = tteconvention_data['result']['name'].replace(" ", "_") + '_volunteer_events.docx'
    doc_save = 'downloads/' + doc_name
    document.save(doc_save)
    return(doc_name)

# -----------------------------------------------------------------------
# Allow only CSV files
# -----------------------------------------------------------------------
def allowed_file(filename):
    extensions=config.ALLOWED_EXTENSIONS

    return '.' in filename and \
        filename.rsplit('.', 1)[1].lower() in extensions

# -----------------------------------------------------------------------
# Convention Functions
# -----------------------------------------------------------------------
# -----------------------------------------------------------------------
# Pull Convention listing from TTE for TRI
# -----------------------------------------------------------------------
def gettteconventions(ttesession):
    params = {'session_id': ttesession['id']}
    response = requests.get(config.tte_url + "/group/" + config.tte_group_id + "/conventions", params=params)
    data = response.json()
    convention_count = 0
    conventions = {}
    for convention in data['result']['items']:
        toadd = {'name': convention['name'], 'id': convention['id']}
        conventions[convention_count]= toadd
        convention_count = convention_count + 1
    return(conventions)

# -----------------------------------------------------------------------
# Pull Convention Data from the TTE API
# -----------------------------------------------------------------------
def tte_convention_api_get(ttesession,tteconvention_id):
    # print ('tte_convention_api_get testing')
    # Call the global so we can modify it in the function
    global tteconvention_data
    convention_info = {}
    # API Pull from TTE to get the convention information and urls needed to process the convention.
    con_params = {'session_id': ttesession['id'], '_include_relationships': 1, '_include': 'description'}
    convention_response = requests.get(config.tte_url + "/convention/" + tteconvention_id, params= con_params)
    tteconvention_data = convention_response.json()
    # API Pull from TTE to get the external json information
    con_jparams = {'session_id': ttesession['id']}
    convention_jresponse = requests.get(config.tte_url + "/convention/" + tteconvention_id + '/external_jsons', params= con_jparams)
    convention_jjson = convention_jresponse.json()
    tteconvention_data['result']['external_jsons'] = convention_jjson['result']['items']
    # API Pull from TTE to get
    event_data = tte_events_api_get(ttesession,tteconvention_id)
    for event in event_data:
        # Get the slots this event is assigned to
        slots_url = 'https://tabletop.events' + event['_relationships']['slots']
        event_slots = tte_event_slots_api_get(ttesession,tteconvention_id,slots_url)
        slot_tables = slots_parse(event_slots)
        event['event_tables'] = slot_tables
        # Get the hosts this event has
        # hosts_url = field['_relationships']['hosts']
        # event_hosts = tte_event_hosts_api_get(ttesession,tteconvention_id,hosts_url)
        # field['event_hosts'] = event_hosts
    # API Pull from TTE to get the volunteer information
    volunteer_data = tte_convention_volunteer_api_get(ttesession,tteconvention_id)
    # Populate dictionary with the info pulled from TTE
    tteconvention_data['result']['geolocation_name'] = tte_geolocation_byid_api_get(ttesession)
    tteconvention_data['result']['days'] = tte_convention_days_api_get(ttesession,tteconvention_id)
    tteconvention_data['events'] = event_data
    tteconvention_data['volunteers'] = volunteer_data
    print (tteconvention_data['volunteers'])
    return()

# -----------------------------------------------------------------------
# Create a new Convention
# -----------------------------------------------------------------------
def tte_convention_convention_api_post(ttesession,new_convention):
    print ('debug tte_convention_convention_api_post')
    # Declarations
    # Function Calls
    geolocation_id = tte_geolocation_api_get(ttesession,new_convention)
    # Define parameters to create the convention
    convention_url = '/api/convention'
    convention_params = {
                        'session_id': ttesession['id'],
                        'website_uri': 'https://theroleinitiative.org',
                        'name': new_convention['name'],
                        'description': new_convention['description'],
                        'facebook_page': 'https://www.facebook.com/theroleinitiative/',
                        'generic_ticket_price': 0,
                        'group_id': config.tte_group_id,
                        'slot_duration': 30,
                        'twitter_handle': '@_roleinitiative',
                        'email_address': 'events@theroleinitiative.org',
                        'phone_number': new_convention['phone_number'],
                        'geolocation_id': geolocation_id,
                        'volunteer_management': 'enabled',
                        '_include_relationships':1,
                        }
    convention_response = requests.post('https://tabletop.events' + convention_url, params= convention_params)
    convention_json = convention_response.json()
    tteconvention_id = convention_json['result']['id']
    # Create each day of the convention
    tte_convention_days_api_post(ttesession,tteconvention_id,new_convention)
        # Create the standard TRI custom form
    convention_conventionjson_params = {
                        'session_id': ttesession['id'],
                        'convention_id': tteconvention_id,
                        'name': 'volunteer_custom_fields'
                        }
    convention_conventionjson_json_data = {
                        'json': [
                        {
                            "required" : "1",
                            "label" : "Emergency Contact: Name, phone number, relationship",
                            "name" : "volunteeremergencycontact",
                            "edit" : "0",
                            "type" : "text",
                            "conditional" : "0",
                            "view" : "1",
                            "sequence_number" : "3"
                         },
                         {
                            "required" : "0",
                            "label" : "Previous Convention/D&D Volunteer experience",
                            "type" : "textarea",
                            "conditional" : "0",
                            "name" : "volunteerexperience",
                            "edit" : "0",
                            "sequence_number" : "2",
                            "view" : "1"
                         },
                         {
                            "view" : "1",
                            "sequence_number" : "4",
                            "edit" : "0",
                            "name" : "volunteerlevel",
                            "type" : "select",
                            "conditional" : "0",
                            "options" : "Hotel\n4 Day\n1 day\n1 slot",
                            "label" : "Volunteer Level - Hotel level requires committing to 24 hours over the 4 days of the convention.  Badge Level requires 12 hours, Day level requires 4 hours.  1 slot is 2 hours.  At this time we cannot confirm Hotels Slots will be available for the convention, but if you are interested in volunteering at that level still select that as an option please.",
                            "required" : "1"
                         },
                         {
                            "required" : "0",
                            "label" : "Shirt Size",
                            "options" : "S\nM\nL\nXL\nXXL\n3X\n4X\n5X",
                            "type" : "select",
                            "conditional" : "0",
                            "name" : "volunteershirtsize",
                            "edit" : "0",
                            "view" : "1",
                            "sequence_number" : "7"
                         },
                         {
                            "label" : "Other comments (accommodations requests, allergies we should be aware of, other things you feel you should share, etc.)",
                            "required" : "0",
                            "type" : "textarea",
                            "conditional" : "0",
                            "edit" : "0",
                            "name" : "volunteerother",
                            "sequence_number" : "11",
                            "view" : "1"
                         },
                         {
                            "sequence_number" : "9",
                            "view" : "1",
                            "conditional" : "0",
                            "type" : "text",
                            "edit" : "0",
                            "name" : "volunteerlocation",
                            "label" : "Where are you coming from (City/State)",
                            "required" : "1"
                         },
                         {
                            "view" : "1",
                            "sequence_number" : "1",
                            "required" : "0",
                            "label" : "What pronouns do you use for yourself?",
                            "name" : "volunteerpronouns",
                            "edit" : "0",
                            "conditional" : "0",
                            "type" : "text"
                         },
                         {
                            "sequence_number" : "8",
                            "view" : "1",
                            "edit" : "0",
                            "name" : "volunteersource",
                            "conditional" : "0",
                            "type" : "text",
                            "label" : "How did you hear about us?",
                            "required" : "1"
                         },
                         {
                            "options" : "None\n1\n2\n3\n4",
                            "required" : "1",
                            "label" : "Tier (What is the highest Tier you are comfortable GMing, enter None if you do not want to GM at all)",
                            "name" : "volunteertiers",
                            "edit" : "0",
                            "conditional" : "0",
                            "type" : "select",
                            "sequence_number" : "6",
                            "view" : "1"
                         },
                         {
                            "sequence_number" : "10",
                            "view" : "1",
                            "conditional_name" : "volunteerlevel",
                            "edit" : "0",
                            "conditional_value" : "Hotel",
                            "name" : "volunteerhotelpref",
                            "conditional" : "1",
                            "type" : "select",
                            "options" : "Male\nFemale\nAny",
                            "label" : "Hotel Rooming Preference",
                            "required" : "0"
                         },
                         {
                            "view" : "1",
                            "sequence_number" : "5",
                            "required" : "1",
                            "label" : "What role are you interested in?  Admin roles are as follows: Runners work with the Admins assigned to the slot, they will help GMs with getting their badges and perform health checks.  Admins will help seat players at tables and check DMs in.  Head admin will be the escalation point for any issues that arise.",
                            "options" : "DM - Adventurers League Only\nDM - Acquisitions Incorporated Only\nDM - Any\nAdmin\nAny",
                            "conditional" : "0",
                            "type" : "select",
                            "name" : "volunteerrole",
                            "edit" : "0"
                         }
                     ]}
    convention_conventionjson_response = requests.post('https://tabletop.events/api/conventionjson', json=convention_conventionjson_json_data, params= convention_conventionjson_params)
    print (convention_conventionjson_response.url)
    convention_conventionjson_json = convention_conventionjson_response.json()
    print (json.dumps(convention_conventionjson_json,indent=2))
    return(tteconvention_id)

# -----------------------------------------------------------------------
# Update a Convention
# -----------------------------------------------------------------------
def tte_convention_convention_api_put(ttesession,update_convention):
    geolocation_id = tte_geolocation_api_get(ttesession,update_convention)
    convention_url = '/api/convention/' + tteconvention_data['result']['id']
    convention_params = {
                        'session_id': ttesession['id'],
                        'name': update_convention['name'],
                        'phone_number': update_convention['phone_number'],
                        'email_address': update_convention['email'],
                        'geolocation_id': geolocation_id,
                        'description': update_convention['description']
                        }
    convention_response = requests.put('https://tabletop.events' + convention_url, params= convention_params)
    convention_json = convention_response.json()
    print (convention_json)
    return()

# -----------------------------------------------------------------------
# Update an existing Convention with TRI standard information
# -----------------------------------------------------------------------
def tte_convention_convention_tristandard_api_put(ttesession):
    print ('debug tte_convention_convention_tristandard_api_put')
    # Declarations
    # Define parameters to update the convention
    convention_url = 'https://tabletop.events/api/convention/' + tteconvention_data['result']['id']
    convention_params = {
                        'session_id': ttesession['id'],
                        'volunteer_custom_fields': [
                            {
                                "required" : "1",
                                "label" : "Emergency Contact: Name, phone number, relationship",
                                "name" : "volunteeremergencycontact",
                                "edit" : "0",
                                "type" : "text",
                                "conditional" : "0",
                                "view" : "1",
                                "sequence_number" : "3"
                             },
                             {
                                "required" : "0",
                                "label" : "Previous Convention/D&D Volunteer experience",
                                "type" : "textarea",
                                "conditional" : "0",
                                "name" : "volunteerexperience",
                                "edit" : "0",
                                "sequence_number" : "2",
                                "view" : "1"
                             },
                             {
                                "view" : "1",
                                "sequence_number" : "4",
                                "edit" : "0",
                                "name" : "volunteerlevel",
                                "type" : "select",
                                "conditional" : "0",
                                "options" : "Hotel\n4 Day\n1 day\n1 slot",
                                "label" : "Volunteer Level - Hotel level requires committing to 24 hours over the 4 days of the convention.  Badge Level requires 12 hours, Day level requires 4 hours.  1 slot is 2 hours.  At this time we cannot confirm Hotels Slots will be available for the convention, but if you are interested in volunteering at that level still select that as an option please.",
                                "required" : "1"
                             },
                             {
                                "required" : "0",
                                "label" : "Shirt Size",
                                "options" : "S\nM\nL\nXL\nXXL\n3X\n4X\n5X",
                                "type" : "select",
                                "conditional" : "0",
                                "name" : "volunteershirtsize",
                                "edit" : "0",
                                "view" : "1",
                                "sequence_number" : "7"
                             },
                             {
                                "label" : "Other comments (accommodations requests, allergies we should be aware of, other things you feel you should share, etc.)",
                                "required" : "0",
                                "type" : "textarea",
                                "conditional" : "0",
                                "edit" : "0",
                                "name" : "volunteerother",
                                "sequence_number" : "11",
                                "view" : "1"
                             },
                             {
                                "sequence_number" : "9",
                                "view" : "1",
                                "conditional" : "0",
                                "type" : "text",
                                "edit" : "0",
                                "name" : "volunteerlocation",
                                "label" : "Where are you coming from (City/State)",
                                "required" : "1"
                             },
                             {
                                "view" : "1",
                                "sequence_number" : "1",
                                "required" : "0",
                                "label" : "What pronouns do you use for yourself?",
                                "name" : "volunteerpronouns",
                                "edit" : "0",
                                "conditional" : "0",
                                "type" : "text"
                             },
                             {
                                "sequence_number" : "8",
                                "view" : "1",
                                "edit" : "0",
                                "name" : "volunteersource",
                                "conditional" : "0",
                                "type" : "text",
                                "label" : "How did you hear about us?",
                                "required" : "1"
                             },
                             {
                                "options" : "None\n1\n2\n3\n4",
                                "required" : "1",
                                "label" : "Tier (What is the highest Tier you are comfortable GMing, enter None if you do not want to GM at all)",
                                "name" : "volunteertiers",
                                "edit" : "0",
                                "conditional" : "0",
                                "type" : "select",
                                "sequence_number" : "6",
                                "view" : "1"
                             },
                             {
                                "sequence_number" : "10",
                                "view" : "1",
                                "conditional_name" : "volunteerlevel",
                                "edit" : "0",
                                "conditional_value" : "Hotel",
                                "name" : "volunteerhotelpref",
                                "conditional" : "1",
                                "type" : "select",
                                "options" : "Male\nFemale\nAny",
                                "label" : "Hotel Rooming Preference",
                                "required" : "0"
                             },
                             {
                                "view" : "1",
                                "sequence_number" : "5",
                                "required" : "1",
                                "label" : "What role are you interested in?  Admin roles are as follows: Runners work with the Admins assigned to the slot, they will help GMs with getting their badges and perform health checks.  Admins will help seat players at tables and check DMs in.  Head admin will be the escalation point for any issues that arise.",
                                "options" : "DM - Adventurers League Only\nDM - Acquisitions Incorporated Only\nDM - Any\nAdmin\nAny",
                                "conditional" : "0",
                                "type" : "select",
                                "name" : "volunteerrole",
                                "edit" : "0"
                             }
                         ]
                        }
    convention_response = requests.put(convention_url, data= convention_params)
    convention_json = convention_response.json()
    print (convention_json)
    return()

# -----------------------------------------------------------------------
# Pull Convention Data from the database
# -----------------------------------------------------------------------
def list_convention_info(tteconvention_id):
    convention = Conventions()
    convention = Conventions.query.filter_by(tteid = tteconvention_id).first()
    convention_data = {}
    try:
        if convention.slots is not None:
            con_slots = json.loads(convention.slots)
            for slot in con_slots:
                try:
                    new_slot = int(slot)
                    convention_data[new_slot] = con_slots[slot]
                except ValueError:
                    pass
        if convention.tables is not None:
            convention_data['tables'] = convention.tables
    except:
        convention_data = None
        pass
    return(convention_data)

# -----------------------------------------------------------------------
# Pull Slots Data from the TTE API for the whole convention that match the time submitt and the event room id
# -----------------------------------------------------------------------
def tte_convention_slots_api_get(ttesession,tteconvention_id,eventslot,event):
    #print ('debug tte_convention_slots_api_get')
    slots_start = 1
    slots_total = 1000
    all_slots = list()
    slots_url = tteconvention_data['result']['_relationships']['slots']
    while slots_total >= slots_start:
        slots_params = {'session_id': ttesession['id'], '_page_number': slots_start, 'daypart_id': eventslot['id'], 'room_id': event['type_room_id']}
        slots_response = requests.get('https://tabletop.events' + slots_url, data= slots_params)
        slots_json = slots_response.json()
        convention_slots = slots_json['result']['items']
        slots_total = int(slots_json['result']['paging']['total_pages'])
        for slots in convention_slots:
            all_slots.append(slots)
        if slots_start < slots_total:
            slots_start = int(slots_json['result']['paging']['next_page_number'])
        elif slots_start == slots_total:
            break
    return(all_slots)

# -----------------------------------------------------------------------
# Pull Slots Data from the TTE API for a specific event
# -----------------------------------------------------------------------
def tte_event_slots_api_get(ttesession,tteconvention_id,slots_url):
    slots_start = 1
    slots_total = 1000
    all_slots = list()
    while slots_total >= slots_start:
        slots_params = {'session_id': ttesession['id'], '_page_number': slots_start}
        slots_response = requests.get(slots_url, params= slots_params)
        slots_json = slots_response.json()
        convention_slots = slots_json['result']['items']
        slots_total = int(slots_json['result']['paging']['total_pages'])
        for slots in convention_slots:
            all_slots.append(slots)
        if slots_start < slots_total:
            slots_start = int(slots_json['result']['paging']['next_page_number'])
        elif slots_start == slots_total:
            break
    return(all_slots)

# -----------------------------------------------------------------------
# Create a list of tables assigned to the slot
# -----------------------------------------------------------------------
def slots_parse(event_slots):
    table_l= []
    slot_tables = []
    for slot in event_slots:
        table_l = slot['name'].split()
        table_str = table_l[0] + ' ' + table_l[1]
        if table_str not in slot_tables:
            slot_tables.append(table_str)
    return (slot_tables)

# -----------------------------------------------------------------------
# Save the event data to CSV
# -----------------------------------------------------------------------
def event_data_csv(events):
    folder = config.DOWNLOAD_FOLDER
    saveloc = '.downloads/eventdata.csv'
    with open(saveloc, mode='w') as csv_file:
        fieldnames = ['event_number', 'name', 'startdaypart_name', 'duration', 'event_tables', 'hosts']
        writer = csv.DictWriter(csv_file, fieldnames=fieldnames,extrasaction='ignore')
        writer.writeheader()
        sorted_events = sorted(events, key = lambda j: j['start_date'])
        for event in sorted_events:
            print (event)
            writer.writerow(event)

    return()

# -----------------------------------------------------------------------
# Save the volunteer data to CSV
# -----------------------------------------------------------------------
def volunteer_data_csv(volunteers):
    folder = config.DOWNLOAD_FOLDER
    saveloc = '.downloads/volunterdata.csv'
    with open(saveloc, mode='w') as csv_file:
        fieldnames = ['email_address', 'firstname', 'lastname', 'shift_list', 'custom_fields']
        writer = csv.DictWriter(csv_file, fieldnames=fieldnames, extrasaction='ignore')
        for volunteer in volunteers:
            shift_list = []
            print (volunteer)
            for shift in volunteer['shifts']:
                shift_list.append(shift['shift_data']['name'])
                volunteer['shift_list'] = shift_list
            writer.writerow(volunteer)
    return()

# -----------------------------------------------------------------------
# Pull Hosts Data from the TTE API for a specific event
# -----------------------------------------------------------------------
def tte_event_hosts_api_get(ttesession,tteconvention_id,hosts_url):
    hosts_start = 1
    hosts_total = 1000
    all_hosts = list()
    while hosts_total >= hosts_start:
        hosts_params = {'session_id': ttesession['id'], '_page_number': hosts_start}
        hosts_response = requests.get('https://tabletop.events' + hosts_url, params= hosts_params)
        hosts_json = hosts_response.json()
        convention_hosts = hosts_json['result']['items']
        hosts_total = int(hosts_json['result']['paging']['total_pages'])
        for hosts in convention_hosts:
            all_hosts.append(hosts)
        if hosts_start < hosts_total:
            hosts_start = int(hosts_json['result']['paging']['next_page_number'])
        elif hosts_start == hosts_total:
            break
    return(all_hosts)

# -----------------------------------------------------------------------
# Get the Geolocation data (for the time zone of the con)
# -----------------------------------------------------------------------
def tte_convention_geolocation_api_get(ttesession,tteconvention_id):
  geolocation_params = {'session_id': ttesession['id'], '_include_related_objects': 'geolocation'}
  geolocation_response = requests.get(config.tte_url + "/convention/" + tteconvention_id, params= geolocation_params)
  geolocation_data = geolocation_response.json()
  geolocation_timezone = geolocation_data['result']['geolocation']['timezone']
  return(geolocation_timezone)

# -----------------------------------------------------------------------
# Parse a File for a convention
# -----------------------------------------------------------------------
def convention_parse(filename,tteconvention_id,tteconvention_name):
    #print('convention_parse')
    # Definitions
    slot = {}
    newheader = []
    convention_slots = []
    convention_tables = []
    convention = {}

    # Open CSV file and verify headers
    with open(filename, newline='') as csvfile:
        reader = csv.DictReader(csvfile)
        for header in reader.fieldnames:
            if 'Slot' in header:
                header_l = header.rsplit()
                newheader.append('slot ' + header_l[1])
            if 'Length' in header:
                newheader.append('length')
            if 'Table Start' in header:
                newheader.append('table_start')
            if 'Table End' in header:
                newheader.append('table_end')
            if 'Table Type' in header:
                newheader.append('table_type')
        reader.fieldnames = newheader
        for room_info in reader:
            tables = {}
            new_slot = {}
            for field in room_info:
                # Create the dict of slot time and length of each slot if the field is a slot field
                if 'slot' in field and room_info[field] is not 'X':
                    slot_num = field.rsplit()
                    new_slot = {'slot': slot_num[1] , 'time': room_info[field], 'length': room_info['length']}
                    # Add to a list of all the slots for the convention
                    convention_slots.append(new_slot)
            # Create a dict for each room of the convention
            tables = {'table_type': room_info['table_type'], 'table_start': room_info['table_start'],'table_end': room_info['table_end']}
            # Add the tables dict to a list of rooms
            convention_tables.append(tables)
        convention['slots'] = convention_slots
        convention['tables'] = convention_tables
        # save_convention(convention,tteconvention_id,tteconvention_name)
        return(convention)

# -----------------------------------------------------------------------
# Save a convention to the database
# -----------------------------------------------------------------------
def save_convention(convention,tteconvention_id,tteconvention_name):
    convention_exist = list_convention_info(tteconvention_id)
    new_convention = Conventions()
    #Check to see if the convention already exists
    #If the convention doesn't, Setup new convention database entry
    if convention_exist is None:
        new_convention.tteid = tteconvention_id
        new_convention.name = tteconvention_name
        new_convention.tables = convention['tables']
        new_convention.slots = convention['slots']
        db.session.merge(new_convention)
    else:
        saved = 'exists'
    try:
        db.session.commit()
        saved = 'saved'
        return (saved)
    except:
        logger.exception("Cannot save new convention")
        db.session.rollback()
        saved = 'failed'
        return (saved)

# -----------------------------------------------------------------------
# Get Convention Volunteer Information
# -----------------------------------------------------------------------
def tte_convention_volunteer_api_get(ttesession,tteconvention_id):
    volunteer_start = 1
    volunteer_total = 1000
    all_volunteers = list()
    ttevolunteer_url = 'https://tabletop.events' + tteconvention_data['result']['_relationships']['volunteers']
    while volunteer_total >= volunteer_start:
        volunteer_params = {'session_id': ttesession['id'], 'convention_id': tteconvention_id, '_page_number': volunteer_start}
        volunteer_response = requests.get(ttevolunteer_url, params= volunteer_params)
        volunteer_json = volunteer_response.json()
        volunteer_data = volunteer_json['result']['items']
        volunteer_total = int(volunteer_json['result']['paging']['total_pages'])
        volunteer_start = int(volunteer_json['result']['paging']['page_number'])
        for volunteer in volunteer_data:
            all_volunteers.append(volunteer)
        if volunteer_start < volunteer_total:
            volunteer_start = int(volunteer_json['result']['paging']['next_page_number'])
        elif volunteer_start == volunteer_total:
            break
        else:
            break
    return(all_volunteers)

# -----------------------------------------------------------------------
# Get days of the convention
# -----------------------------------------------------------------------
def tte_convention_days_api_get(ttesession,tteconvention_id):
    #Declarations
    day_info = []
    # Get the data on the convention
    tteconvention_days_url = 'https://tabletop.events' + tteconvention_data['result']['_relationships']['days']
    # Use the day url to get data on the days
    day_params = {'session_id': ttesession, 'convention_id': tteconvention_id}
    day_response = requests.get(tteconvention_days_url, params= day_params)
    day_data = day_response.json()
    # Create a datetime value for each day, add to a new dict
    for item in day_data['result']['items']:
         dt = datetime.datetime.strptime(item['start_date'], '%Y-%m-%d %H:%M:%S')
         et = datetime.datetime.strptime(item['end_date'], '%Y-%m-%d %H:%M:%S')
         day_info.append({'id' : item['id'], 'day_time' : dt, 'end_time': et})
    return(day_info)

# -----------------------------------------------------------------------
# Post the Convention Days
# -----------------------------------------------------------------------
def tte_convention_days_api_post(ttesession,tteconvention_id,new_convention):
    # Declarations
    all_dates = new_convention['dates'].split('\r\n')
    tteconvention_days_url = 'https://tabletop.events/api/conventionday'
    for date in all_dates:
        day_date = date + ' 12:00:00 AM'
        start_date = datetime.datetime.strptime(day_date, '%m/%d/%Y %I:%M:%S %p')
        start_date_utc = datetime_utc_convert(ttesession,tteconvention_id,start_date)
        start_day = start_date_utc.strftime('%Y-%m-%d %H:%M:%S')
        end_date_utc = start_date_utc + datetime.timedelta(days=1)
        end_day = end_date_utc.strftime('%Y-%m-%d %H:%M:%S')
        day_name = start_date.strftime('%a %b %d')
        day_params = {
            'session_id': ttesession['id'],
            'attendee_start_date': start_day,
            'attendee_end_date': end_day,
            'start_date': start_day,
            'end_date': end_day,
            'convention_id': tteconvention_id,
            'name': day_name,
            'day_type': 'events'
        }
        day_response = requests.post(tteconvention_days_url, params= day_params)
        day_json = day_response.json()
    return()

# -----------------------------------------------------------------------
# Get the id for day parts
# -----------------------------------------------------------------------
def tte_convention_dayparts_api_get(ttesession,tteconvention_id):
    day_parts_start = 1
    day_parts_total = 1000
    all_dayparts = list()
    dayparts_url = tteconvention_data['result']['_relationships']['dayparts']

    while day_parts_total >= day_parts_start:
        dayparts_params = {'session_id': ttesession['id'], '_page_number': day_parts_start}
        dayparts_response = requests.get('https://tabletop.events' + dayparts_url, params= dayparts_params)
        dayparts_data = dayparts_response.json()
        convention_dayparts = dayparts_data['result']['items']
        day_parts_total = int(dayparts_data['result']['paging']['total_pages'])
        for dayparts in convention_dayparts:
            dayparts['datetime'] = datetime.datetime.strptime(dayparts['start_date'],'%Y-%m-%d %H:%M:%S')
            all_dayparts.append(dayparts)
        if day_parts_start < day_parts_total:
            day_parts_start = int(dayparts_data['result']['paging']['next_page_number'])
        elif day_parts_start == day_parts_total:
            break
    return(all_dayparts)

# -----------------------------------------------------------------------
# Delete all dayparts from TTE
# -----------------------------------------------------------------------
def tte_convention_dayparts_api_delete(ttesession,tteconvention_id,all_dayparts):
    for daypart in all_dayparts:
        daypart_delete_params = {'session_id': ttesession['id']}
        daypart_delete_url = 'https://tabletop.events/api/daypart/' + daypart['id']
        daypart_delete_response = requests.delete(daypart_delete_url, params= daypart_delete_params)
        daypart_delete_data = daypart_delete_response.json()
    return()

# -----------------------------------------------------------------------
# Delete a volunteer from TTE for a specific convention
# -----------------------------------------------------------------------
def tte_convention_volunteer_api_delete(ttesession,tteconvention_id,volunteer_id):
    volunteer_delete_params = {'session_id': ttesession['id']}
    volunteer_delete_url = 'https://tabletop.events/api/volunteer/' + volunteer_id
    print (volunteer_delete_url)
    volunteer_delete_response = requests.delete(volunteer_delete_url, params= volunteer_delete_params)
    volunteer_delete_data = volunteer_delete_response.json()
    print (volunteer_delete_data)
    return()

# -----------------------------------------------------------------------
# Volunteer Functions
# -----------------------------------------------------------------------
# -----------------------------------------------------------------------
# Ingest Volunteers
# -----------------------------------------------------------------------
def volunteer_parse(filename,tteconvention_id):
    # Definitions
    newheader = []

    # Open CSV file and verify headers
    with open(filename, newline='') as csvfile:
        reader = csv.DictReader(csvfile)
        for header in reader.fieldnames:
            if 'Email Address' in header:
                newheader.append('email')
            elif 'Name' in header:
                newheader.append('name')
            elif 'Role' in header:
                newheader.append('role')
            elif 'Hours' in header:
                newheader.append('hours')
            elif 'Tiers' in header:
                newheader.append('tiers')
            elif 'Slot' in header:
                header_l = header.rsplit()
                newheader.append('slot ' + header_l[1])
        reader.fieldnames = newheader
        for vol in reader:
            saved = volunteer_save(vol,tteconvention_id)
        return(saved)

# -----------------------------------------------------------------------
# Query if user exists in TTE
# -----------------------------------------------------------------------
def tte_user_api_pull(ttesession,volunteer_email):
    print ('tte_user_api_pull')
    volunteer_params = {'session_id': ttesession['id']}
    volunteer_url = 'https://tabletop.events' + '/api/user' + '?query=' + volunteer_email
    volunteer_response = requests.get(volunteer_url, params= volunteer_params)
    volunteer_json = volunteer_response.json()
    try:
        for vol in volunteer_json['result']['items']:
            volunteer_id = vol['id']
            print (vol['real_name'],vol['id'])
    except:
        print ('Unable to find: ', volunteer_email)
        volunteer_id = None
    return(volunteer_id)

# -----------------------------------------------------------------------
# Add user to TTE
# -----------------------------------------------------------------------
def tte_user_add(ttesession,volunteer_email,volunteer_name,tteconvention_id):
    #  print ('tte_user_add')
    volunteer_full_name = volunteer_name.rsplit()
    volunteer_first = volunteer_full_name[0]
    volunteer_last = volunteer_full_name[1]
    useradd_params = {'session_id': ttesession['id'],'convention_id' : tteconvention_id,'email_address' : volunteer_email,'firstname' : volunteer_first,'lastname' : volunteer_last,'phone_number' : '555-555-5555'}
    volunteer_response = requests.post('https://tabletop.events/api/volunteer/by-organizer', params= useradd_params)
    volunteer_data = volunteer_response.json()
    try:
        volunteer_id = volunteer_data['result']['id']
        return(volunteer_id)
    except:
        print ('Unable to add: ', volunteer_email)
        return()


# -----------------------------------------------------------------------
# Get Volunteer Shift Information
# -----------------------------------------------------------------------
def tte_volunteer_shifts_api_get(ttesession,tteconvention_id,volunteer_id):
    volunteer_shifts_start = 1
    volunteer_shifts_total = 1000
    all_volunteer_shifts = list()
    ttevolunteer_shifts_url = 'https://tabletop.events/api/convention/' + tteconvention_id + '/volunteershifts'
    while volunteer_shifts_total >= volunteer_shifts_start:
        volunteer_shifts_params = {'session_id': ttesession, 'convention_id': tteconvention_id, 'volunteer_id':volunteer_id , '_page_number': volunteer_shifts_start}
        volunteer_shifts_response = requests.get(ttevolunteer_shifts_url, params= volunteer_shifts_params)
        volunteer_shifts_json = volunteer_shifts_response.json()
        volunteer_shifts_data = volunteer_shifts_json['result']['items']
        volunteer_shifts_total = int(volunteer_shifts_json['result']['paging']['total_pages'])
        volunteer_shifts_start = int(volunteer_shifts_json['result']['paging']['page_number'])
        for volunteer_shift in volunteer_shifts_data:
            volunteer_shift['shift_data'] = tte_shift_api_get(ttesession,tteconvention_id,volunteer_shift['shift_id'])
            all_volunteer_shifts.append(volunteer_shift)
        if volunteer_shifts_start < volunteer_shifts_total:
            volunteer_shifts_start = int(volunteer_shifts_json['result']['paging']['next_page_number'])
        elif volunteer_shifts_start == volunteer_shifts_total:
            break
        else:
            break
    return(all_volunteer_shifts)

# -----------------------------------------------------------------------
# Get Shift Information
# -----------------------------------------------------------------------
def tte_shift_api_get(ttesession,tteconvention_id,shift_id):
    shift_start = 1
    shift_total = 1000
    all_shift = list()
    tteshift_url = 'https://tabletop.events/api/shift/' + shift_id
    shift_params = {'session_id': ttesession, 'convention_id': tteconvention_id, '_page_number': shift_start}
    shift_response = requests.get(tteshift_url, params= shift_params)
    shift_json = shift_response.json()
    shift_data = shift_json['result']
    return(shift_data)

# -----------------------------------------------------------------------
# Slot Functions
# -----------------------------------------------------------------------
# -----------------------------------------------------------------------
# Delete shofts from database
# -----------------------------------------------------------------------
def database_slot_delete(tteconvention_id):
    convention = Conventions()
    convention = Conventions.query.filter_by(tteid=tteconvention_id).first()
    convention.slots = None
    try:
        db.session.commit()
        deleted = 'Deleted all slots'
        return (deleted)
    except:
        logger.exception("Cannot delete slots")
        db.session.rollback()
        deleted = 'failed'
        return (deleted)

# -----------------------------------------------------------------------
# Post to TTE the Volunteer Shifts
# -----------------------------------------------------------------------
def tte_convention_volunteer_shift_api_post(ttesession,tteconvention_id,convention_info):
    print ('tte_convention_volunteer_shift_api_post')
    # Get the information on Convention Days
    day_info = tte_convention_days_api_get(ttesession,tteconvention_id)
    # Get the information on preexisting shifts
    shiftypes_info = tte_convention_volunteer_shifttypes_api_get(ttesession,tteconvention_id)
    # If there are any existing shifts, look for a match on the name
    if len(shiftypes_info) != 0:
        for shifttype in shiftypes_info:
            if 'Slot' in shifttype['name']:
                shifttype_name = shifttype['name']
                shifttype_id = shifttype['id']
                print ('Found shift type ', shifttype_name,shifttype_id)
            else:
                pass
    else:
        shifttype_name = 'Slot'
        shifttype_id = tte_convention_volunteer_shifttypes_api_post(ttesession,tteconvention_id,shifttype_name)
        print ('Created shift type ', shifttype_name,shifttype_id)
    # For each slot, get the information we need to be able to post the a volunteer shift
    for slot in convention_info['slots']:
        shift_name = 'Slot ' + slot['slot']
        slot_length = int(slot['length'])
        shift_actual = datetime.datetime.strptime(slot['time'], '%m/%d/%y %I:%M:%S %p')
        shift_start = datetime_utc_convert(ttesession,tteconvention_id,shift_actual)
        shift_end = shift_start + datetime.timedelta(hours=slot_length)
        for day in day_info:
            slot_date = datetime.date(shift_actual.year,shift_actual.month,shift_actual.day)
            shift_date = datetime.date(day['day_time'].year,day['day_time'].month,day['day_time'].day)
            print (slot_date,shift_date)
            # Compare the dates of the slot and the shift to get the tteid to use to post the shift
            if slot_date == shift_date:
                day_id = day['id']
                shift_params = {
                'session_id': ttesession['id'],
                'convention_id': tteconvention_id,
                'name': shift_name,
                'quantity_of_volunteers': '255',
                'start_time': shift_start,
                'end_time': shift_end,
                'conventionday_id': day_id,
                'shifttype_id': shifttype_id
                }
                shift_response = requests.post(config.tte_url + '/shift', params= shift_params)
                shift_json = shift_response.json()
                print (shift_json)
    return()

# -----------------------------------------------------------------------
# Pull shifts from TTE
# -----------------------------------------------------------------------
def tte_convention_volunteer_shifttypes_api_get(ttesession,tteconvention_id):
    shifttypes_url = 'https://tabletop.events' + tteconvention_data['result']['_relationships']['shifttypes']
    shifttypes_params = {'session_id': ttesession, 'convention_id': tteconvention_id}
    shifttypes_response = requests.get(shifttypes_url, params= shifttypes_params)
    shifttypes_json = shifttypes_response.json()
    shifttypes_data = shifttypes_json['result']['items']
    return (shifttypes_data)

# -----------------------------------------------------------------------
# Post shiftstypes to TTE
# -----------------------------------------------------------------------
def tte_convention_volunteer_shifttypes_api_post(ttesession,tteconvention_id,shifttype_name):
    shifttypes_params = {'session_id': ttesession['id'], 'convention_id': tteconvention_id, 'name': shifttype_name}
    shifttypes_response = requests.post(config.tte_url + '/shifttype', params= shifttypes_params)
    shifttypes_json = shifttypes_response.json()
    shifttypes_id = shifttypes_json['result']['id']
    return(shifttypes_id)

# -----------------------------------------------------------------------
# Post slots to TTE as Day Parts
# -----------------------------------------------------------------------
def tte_convention_dayparts_api_post(ttesession,tteconvention_id,convention_info):
    #Declarations
    slots = {}
    # Get data on the days
    day_info = tte_convention_days_api_get(ttesession,tteconvention_id)
    # Loop through the day in 30 minute increments
    for day in day_info:
        day_id = day['id']
        day_start = day['day_time']
        day_end = day['end_time']
        daypart_time = day_start
        while daypart_time < day_end:
            convention_datetime = datetime_timezone_convert(ttesession,tteconvention_id,daypart_time)
            daypart_name = datetime.datetime.strftime(convention_datetime, '%a %I:%M %p')
            # API Post to TTE (Day Parts)
            daypart_params = {'session_id': ttesession['id'], 'convention_id': tteconvention_id, 'name': daypart_name, 'start_date': daypart_time, 'conventionday_id': day_id}
            daypart_response = requests.post(config.tte_url + '/daypart', params= daypart_params)
            daypart_data = daypart_response.json()
            daypart_time = daypart_time + datetime.timedelta(minutes= 30)
    return('saved')

# -----------------------------------------------------------------------
# Pull TTE Volunteer Shifts
# -----------------------------------------------------------------------
def tte_convention_volunteer_shifts_api_get(ttesession,tteconvention_id):
    shifts_start = 1
    shifts_total = 1000
    all_shifts = list()
    # Get the data on the convention
    tteconvention_shifts_url = 'https://tabletop.events' + tteconvention_data['result']['_relationships']['shifts']
    while shifts_total >= shifts_start:
        shifts_params = {'session_id': ttesession['id'], '_include_related_objects': 'shifttype'}
        shifts_response = requests.get(tteconvention_shifts_url, params= shifts_params)
        shifts_json = shifts_response.json()
        shifts_total = int(shifts_json['result']['paging']['total_pages'])
        shifts_data = shifts_json['result']['items']
        for shifts in shifts_data:
            all_shifts.append(shifts)
        if shifts_start < shifts_total:
            shifts_start = int(shifts_json['result']['paging']['next_page_number'])
        elif shifts_start == shifts_total:
            break
    return(all_shifts)

# -----------------------------------------------------------------------
# Delete all shifts from TTE
# -----------------------------------------------------------------------
def tte_convention_volunteer_shifts_api_delete(ttesession,tteconvention_id,all_shifts):
    for shift in all_shifts:
        shift_delete_params = {'session_id': ttesession['id']}
        shift_delete_url = 'https://tabletop.events/api/shift/' + shift['id']
        shift_delete_response = requests.delete(shift_delete_url, params= shift_delete_params)
        shift_delete_data = shift_delete_response.json()
    return()

# -----------------------------------------------------------------------
# Event Functions
# -----------------------------------------------------------------------
# -----------------------------------------------------------------------
# Parse Events Matrix
# -----------------------------------------------------------------------
def event_parse(filename,tteconvention_id,tteconvention_name):
    #Definitions
    newheader = []
    savedevents = []
    # Open CSV file and verify headers
    with open(filename, newline='') as csvfile:
        reader = csv.DictReader(csvfile)
        for header in reader.fieldnames:
            if 'Event Name' in header:
                newheader.append('name')
            elif "Datetime" in header:
                newheader.append('datetime')
            elif 'Duration' in header:
                newheader.append('duration')
            elif 'Table Count' in header:
                newheader.append('tablecount')
            elif 'Hosts' in header:
                newheader.append('hosts')
            elif 'Type' in header:
                newheader.append('type')
            elif 'Tier' in header:
                newheader.append('tier')
        reader.fieldnames = newheader
        for event in reader:
            event_hosts = []
            event_hosts = event['hosts'].split()
            event['hosts'] = event_hosts
            savedevents.append(event)
        return(savedevents)

# -----------------------------------------------------------------------
# Push Events to TTE
# -----------------------------------------------------------------------
def tte_convention_events_api_post(ttesession,tteconvention_id,savedevents):
    # print ('tte_convention_events_api_post testing')
    all_events = []

    # Function to create event type and event room type
    def add_event_type(ttesession,tteconvention_id,event):
        # Create the event type
        event['type_id'] = tte_convention_events_type_api_post(ttesession,tteconvention_id,event)
        # Assign the room that matches the event type and return that id.
        event['type_room_id'] = tte_convention_event_type_room_api_post(ttesession,tteconvention_id,event)
        return(event)

    # For each event, gather the information needed to post the event
    # Get the convention days information
    convention_days = tte_convention_days_api_get(ttesession,tteconvention_id)
    # Get the dayparts for the convention
    convention_dayparts = tte_convention_dayparts_api_get(ttesession,tteconvention_id)
    for event in savedevents:
        event_type_l = []
        print (event)
        #Get the event types from TTE
        event_types = tte_convention_eventtypes_api_get(ttesession,tteconvention_id)
        # Compare the Name of the event types (if any exist) with the provided type listed for the event
        event_type_l = [type for type in event_types if type['name'] == event['type']]
        # If there are event types and a match is found, assign the id of the match to the event'
        if len(event_type_l) !=0:
            all_rooms = tte_convention_rooms_api_get(ttesession,tteconvention_id)
            for e in event_type_l:
                if e['name'] == event['type']:
                    event['type_id'] = e['id']
                    for room in all_rooms:
                        if event['type'] == room['name']:
                            event['type_room_id'] = room['id']
                else:
                    event = add_event_type(ttesession,tteconvention_id,event)
        # If no event types exist, create a new Event Type and return the TTE id for that Type, and create an Event Room Type ID.
        else:
            if event['tier'] !='':
                print ('Adding Event Type to TTE: ', event['type'], ' Tier: ', event['tier'])
            else:
                print ('Adding Event Type to TTE: ', event['type'])
            event = add_event_type(ttesession,tteconvention_id,event)
        # Calculate the datetime value of the event
        event['duration'] = int(event['duration'])
        event['unconverted_datetime'] = datetime.datetime.strptime(event['datetime'],'%m/%d/%y %I:%M:%S %p')
        #Convert the datetime value to UTC
        event['datetime_utc'] = datetime_utc_convert(ttesession,tteconvention_id,event['unconverted_datetime'])
        # Identify the Day Id for the convention
        for day in convention_days:
            day['date_check'] = datetime.date(day['day_time'].year,day['day_time'].month,day['day_time'].day)
            event['date_check'] = datetime.date(event['datetime_utc'].year,event['datetime_utc'].month,event['datetime_utc'].day)
            if event['date_check'] == day['date_check']:
                event['day_id'] = day['id']
        # Define a list to be filled with the slot times used (in increments of 30 minutes) in the event
        all_event_times = []
        for x in range(0,event['duration'],30):
            event_time = event['datetime_utc'] + datetime.timedelta(minutes=x)
            all_event_times.append(event_time)
        # Parse through the convention dayparts and the times of the event
        # Compare to see if there are matches to determine the TTE ID of the times of the event
        # Define a list to be filled with the ids and datetimes of the times of the event
        event_time_info = []
        for event_time in all_event_times:
            for dayparts in convention_dayparts:
                daypart_event_time = {}
                # Find the id of the daypart for the start of the event and add that to the event dict
                # Add to the list of slot times and ids
                if dayparts['datetime'] == event_time and event['datetime_utc'] == dayparts['datetime']:
                    daypart_event_time['id'] = dayparts['id']
                    daypart_event_time['datetime'] = dayparts['datetime']
                    event_time_info.append(daypart_event_time)
                    event['dayparts_start_id'] = dayparts['id']
                # Add other ids of correspdonging slot times that fall within the event
                elif dayparts['datetime'] == event_time and event['datetime_utc'] != dayparts['datetime']:
                    daypart_event_time['id'] = dayparts['id']
                    daypart_event_time['datetime'] = dayparts['datetime']
                    event_time_info.append(daypart_event_time)
        # Verify an event has a ID for the day, ID for the Event Type, and ID for the Day Part
        if event['day_id'] and event['type_id'] and event['dayparts_start_id']:
            # Create the Event
            event_data = tte_event_api_post(ttesession,tteconvention_id,event)
            print ('Added new Event to TTE: ', event_data['name'], event['unconverted_datetime'], event_data['id'])
            event['id'] = event_data['id']
            # Add slots for the event (assigns tables and times) as many times as there are tables for the event
            for i in range(0,int(event['tablecount']),1):
                convention_slots_info = []
                for eventslot in event_time_info:
                    # Get the slots for the convention that span the daypart_event_time, and the event room id
                    convention_slots = tte_convention_slots_api_get(ttesession,tteconvention_id,eventslot,event)
                    convention_slots_info.extend(convention_slots)
                # Create a list of slots that are at the same space (table) and are available
                old_space = None
                event_slot_list = []
                for slot in convention_slots_info:
                    print (slot['name'], slot['is_assigned'])
                    if old_space == None and slot['is_assigned'] == 0:
                        old_space = slot
                        event_slot_list.append(slot)
                    if old_space == None and slot['is_assigned'] != 0:
                        pass
                    elif old_space['space_id'] == slot['space_id'] and slot['is_assigned'] == 0:
                        event_slot_list.append(slot)
                    else:
                        pass
                print (event_slot_list)
                # Schedule each slot
                all_event_slots = []
                for conslot in event_slot_list:
                    event_slot_url = 'https://tabletop.events/api/slot/' + conslot['id']
                    event_slot_params = {'session_id': ttesession['id'], 'event_id': event['id']}
                    event_slot_response = requests.put(event_slot_url, params=event_slot_params)
                    event_slot_json = event_slot_response.json()
                    # print (event_slot_json)
                    try:
                        all_event_slots.append(event_slot_json['result']['id'])
                        event['slots'] = all_event_slots
                        print ('Added event to slot ', event_slot_json['result']['name'])
                    except:
                        print ('Unable to add slot', event_slot_json['result']['name'])
                # If a host email address is included for the event, add the host to the event
                if event['hosts'] != '':
                    for host in event['hosts']:
                        try:
                            volunteer = next((vol for vol in tteconvention_data['volunteers'] if vol['email_address'] == host), None)
                            host = tte_event_host_post(ttesession,event_data['id'],volunteer['user_id'])
                            print (volunteer['name'], 'added to', event['name'])
                        except:
                            print ('Unable to add host to event', event['name'])
            all_events.append(event)
    return(all_events)

# -----------------------------------------------------------------------
# Get Event Types Information
# -----------------------------------------------------------------------
def tte_convention_eventtypes_api_get(ttesession,tteconvention_id):
      eventtypes_start = 1
      eventtypes_total = 1000
      all_eventtypes = list()
      # Get the data on the convention
      tteconvention_eventtypes_url = 'https://tabletop.events' + tteconvention_data['result']['_relationships']['eventtypes']
      # Loop through the eventtypes for the convention
      while eventtypes_total >= eventtypes_start:
        eventtypes_params = {'session_id': ttesession, 'convention_id': tteconvention_id, '_include': 'custom_fields'}
        eventtypes_response = requests.get(tteconvention_eventtypes_url, params= eventtypes_params)
        eventtypes_json = eventtypes_response.json()
        eventtypes_data = eventtypes_json['result']['items']
        eventtypes_total = int(eventtypes_json['result']['paging']['total_pages'])
        for eventtypes in eventtypes_data:
            all_eventtypes.append(eventtypes)
        if eventtypes_start < eventtypes_total:
            eventtypes_start = int(eventtypes_json['result']['paging']['next_page_number'])
        elif eventtypes_start == eventtypes_total:
            break
      return(all_eventtypes)

# -----------------------------------------------------------------------
# Create a new Event
# -----------------------------------------------------------------------
def tte_event_api_post(ttesession,tteconvention_id,event):
    print ('testing tte_event_api_post')
    if event['tier'] != '':
        event_tier = {
        'custom_fields':
        {
        'tier': event['tier'],
        'view' : 1
        }
        }
        event_params = {'session_id': ttesession['id'], 'convention_id': tteconvention_id, 'name' : event['name'], 'max_tickets' : 6, 'priority' : 3, 'age_range': 'all', 'type_id' : event['type_id'], 'conventionday_id': event['day_id'], 'duration' : event['duration'], 'alternatedaypart_id' : event['dayparts_start_id'], 'preferreddaypart_id' : event['dayparts_start_id']}
        event_response = requests.post('https://tabletop.events/api/event', json=event_tier, params= event_params)
        event_json = event_response.json()
        event_data = event_json['result']
    else:
        event_params = {'session_id': ttesession['id'], 'convention_id': tteconvention_id, 'name' : event['name'], 'max_tickets' : 6, 'priority' : 3, 'age_range': 'all', 'type_id' : event['type_id'], 'conventionday_id': event['day_id'], 'duration' : event['duration'], 'alternatedaypart_id' : event['dayparts_start_id'], 'preferreddaypart_id' : event['dayparts_start_id']}
        event_response = requests.post('https://tabletop.events/api/event', params= event_params)
        event_json = event_response.json()
        event_data = event_json['result']
    return(event_data)

# -----------------------------------------------------------------------
# Post a new Event Type
# -----------------------------------------------------------------------
def tte_convention_events_type_api_post(ttesession,tteconvention_id,event_type):
    #print ('tte_convention_events_type_api_post')
    if event_type['tier'] != '':
        custom_tier = {
        'custom_fields': [{
        'required': '1',
        'type': 'text',
        'label': 'Tier',
        'name': 'tier',
        'conditional': '0',
        'edit': '0',
        'view': '1',
        'sequence_number': '1',
        }]
        }
        events_type_params = {'session_id': ttesession['id'], 'convention_id': tteconvention_id, 'name': event_type['type'], 'limit_volunteers': 0, 'max_tickets': 6, 'user_submittable': 0, 'default_cost_per_slot': 0, 'limit_ticket_availability': 0}
        events_type_response = requests.post(config.tte_url + '/eventtype', json=custom_tier, params= events_type_params)
        events_type_json = events_type_response.json()
        events_type_id = events_type_json['result']['id']
    else:
        events_type_params = {'session_id': ttesession['id'], 'convention_id': tteconvention_id, 'name': event_type['type'], 'limit_volunteers': 0, 'max_tickets': 6, 'user_submittable': 0, 'default_cost_per_slot': 0, 'limit_ticket_availability': 0}
        events_type_response = requests.post(config.tte_url + '/eventtype', params= events_type_params)
        events_type_json = events_type_response.json()
        events_type_id = events_type_json['result']['id']
    return(events_type_id)

# -----------------------------------------------------------------------
# Add an event type room with the allowed event type
# -----------------------------------------------------------------------
def tte_convention_event_type_room_api_post(ttesession,tteconvention_id,event):
    all_rooms = tte_convention_rooms_api_get(ttesession,tteconvention_id)
    for room in all_rooms:
        if event['type'] == room['name']:
            event['type_room_id'] = room['id']
    event_type_room_params = {'session_id': ttesession['id'], 'convention_id': tteconvention_id, 'eventtype_id': event['type_id'], 'room_id': event['type_room_id']}
    event_type_room_response = requests.post(config.tte_url + '/eventtyperoom', params = event_type_room_params)
    event_type_room_id_json = event_type_room_response.json()
    event_type_room_id = event_type_room_id_json['result']['id']
    return(event_type_room_id)

# -----------------------------------------------------------------------
# Add host to an event
# -----------------------------------------------------------------------
def tte_event_host_post(ttesession,event_id,host_id):
  host_params = {'session_id': ttesession['id'] }
  host_url = 'https://tabletop.events/api/event/' + event_id + '/host/' + host_id
  host_response = requests.post(host_url, params= host_params)
  host_json = host_response.json()
  host_data = host_json['result']
  return(host_data)

# -----------------------------------------------------------------------
# Delete all Events for the Convention
# -----------------------------------------------------------------------
def tte_convention_events_api_delete(ttesession,tteconvention_id,allevents):
    for event in allevents:
        event_delete_params = {'session_id': ttesession['id']}
        event_delete_url = 'https://tabletop.events/api/event/' + event['id']
        event_delete_response = requests.delete(event_delete_url, params= event_delete_params)
        event_delete_data = event_delete_response.json()
    return()

# -----------------------------------------------------------------------
# Rooms and Spaces (Tables) Functions
# -----------------------------------------------------------------------
# -----------------------------------------------------------------------
# Post Tables and Rooms to Convention
# -----------------------------------------------------------------------
def tte_convention_roomnsandspaces_api_post(ttesession,tteconvention_id,convention_info):
    print ('tte_convention_roomnsandspaces_api_post:')
    all_spaces = []
    spaces_data = {}
    for room in convention_info['tables']:
        rooms_params = {'session_id': ttesession['id'], 'convention_id': tteconvention_id, 'name': room['table_type'] }
        rooms_response = requests.post(config.tte_url + '/room', params= rooms_params)
        rooms_json = rooms_response.json()
        room_id = rooms_json['result']['id']
        room_name = rooms_json['result']['name']
        for table_num in range(int(room['table_start']),int(room['table_end'])):
            table_name = 'Table ' + str(table_num) + " " + room['table_type']
            spaces_params = {'session_id': ttesession['id'], 'convention_id': tteconvention_id, 'room_id': room_id, 'name': table_name, 'max_tickets': 6}
            spaces_response = requests.post(config.tte_url + '/space', params= spaces_params)
            spaces_json = spaces_response.json()
            space_id = spaces_json['result']['id']
            space_name = spaces_json['result']['name']
            spaces_data = {'space_id': space_id,'space_name': space_name, 'table_type': room_name,'room_id': room_id}
            all_spaces.append(spaces_data)
    return(all_spaces)

# -----------------------------------------------------------------------
# Delete Tables and Rooms in Convention
# -----------------------------------------------------------------------
def tte_convention_roomnsandspaces_api_delete(ttesession,tteconvention_id,tterooms,ttespace):
    # print ('tte_convention_roomnsandspaces_api_delete:')
    for space in ttespace:
        space_delete_params = {'session_id': ttesession['id']}
        space_delete_url = 'https://tabletop.events/api/space/' + space['id']
        space_delete_response = requests.delete(space_delete_url, params= space_delete_params)
        space_delete_data = space_delete_response.json()
    for room in tterooms:
        room_delete_params = {'session_id': ttesession['id']}
        room_delete_url = 'https://tabletop.events/api/room/' + room['id']
        room_delete_response = requests.delete(room_delete_url, params= room_delete_params)
        room_delete_data = room_delete_response.json()
    return()

# -----------------------------------------------------------------------
# Get Table Information
# -----------------------------------------------------------------------
def tte_convention_spaces_api_get(ttesession,tteconvention_id):
    # print ('tte_convention_spaces_api_get:')
    spaces_start = 1
    spaces_total = 1000
    all_spaces = list()
    tteconvention_spaces_url = 'https://tabletop.events' + tteconvention_data['result']['_relationships']['spaces']
    # Loop through the spaces for the convention
    while spaces_total >= spaces_start:
        spaces_params = {'session_id': ttesession, 'convention_id': tteconvention_id, '_page_number': spaces_start}
        spaces_response = requests.get(tteconvention_spaces_url, params= spaces_params)
        spaces_data = spaces_response.json()
        convention_spaces = spaces_data['result']['items']
        spaces_total = int(spaces_data['result']['paging']['total_pages'])
        spaces_start = int(spaces_data['result']['paging']['page_number'])
        for spaces in convention_spaces:
            all_spaces.append(spaces)
        if spaces_start < spaces_total:
            spaces_start = int(spaces_data['result']['paging']['next_page_number'])
        elif spaces_start == spaces_total:
            break
        else:
            break
    return(all_spaces)

# -----------------------------------------------------------------------
# Get Room Information
# -----------------------------------------------------------------------
def tte_convention_rooms_api_get(ttesession,tteconvention_id):
      rooms_start = 1
      rooms_total = 1000
      all_rooms = list()
      tteconvention_rooms_url = 'https://tabletop.events' + tteconvention_data['result']['_relationships']['rooms']
      # Loop through the rooms for the convention
      while rooms_total >= rooms_start:
        rooms_params = {'session_id': ttesession, 'convention_id': tteconvention_id, '_page_number': rooms_start}
        rooms_response = requests.get(tteconvention_rooms_url, params= rooms_params)
        rooms_data = rooms_response.json()
        convention_rooms = rooms_data['result']['items']
        rooms_total = int(rooms_data['result']['paging']['total_pages'])
        rooms_start = int(rooms_data['result']['paging']['page_number'])
        for rooms in convention_rooms:
            all_rooms.append(rooms)
        if rooms_start < rooms_total:
            rooms_start = int(rooms_data['result']['paging']['next_page_number'])
        elif rooms_start == rooms_total and rooms_total:
            break
        else:
          break
      return(all_rooms)

# -----------------------------------------------------------------------
# Get all events for Convention
# -----------------------------------------------------------------------
def tte_events_api_get(ttesession,tteconvention_id):
    events_start = 1
    events_total = 1000
    all_events = list()
    while events_total >= events_start:
        events_url = tteconvention_data['result']['_relationships']['events']
        events_params = {'session_id': ttesession['id'], 'tteconvention_id': tteconvention_id, '_page_number': events_start, '_include_relationships':1, "_include": 'custom_fields', "_include":'hosts'}
        events_response = requests.get('https://tabletop.events' + events_url,params= events_params)
        events_data = events_response.json()
        convention_events = events_data['result']['items']
        events_total = int(events_data['result']['paging']['total_pages'])
        for events in convention_events:
            all_events.append(events)
        if events_start < events_total:
            events_start = int(events_data['result']['paging']['next_page_number'])
        elif events_start == events_total:
            break
    return(all_events)

# -----------------------------------------------------------------------
# Geolocation Functions
# -----------------------------------------------------------------------
# -----------------------------------------------------------------------
# Query for a location id
# -----------------------------------------------------------------------
def tte_geolocation_api_get(ttesession,convention_info):
    geolocation_name = convention_info['location']
    geolocation_url = 'https://tabletop.events/api/geolocation' + '?query=' + convention_info['location']
    geolocation_params = {'session_id': ttesession['id']}
    geolocation_response = requests.get(geolocation_url, params= geolocation_params)
    geolocation_json = geolocation_response.json()
    try:
        for location in geolocation_json['result']['items']:
            new_date = datetime.datetime.strptime(location['date_created'],'%Y-%m-%d %H:%M:%S')
            try:
                old_date
            except NameError:
                old_date = None
            if old_date is None:
                old_date = new_date
                geolocation_id = location['id']
            if new_date > old_date:
                old_date = new_date
                geolocation_id = location['id']
            else:
                geolocation_id = location['id']
                pass
    except:
        print ('Could not find location', convention_info, 'adding to TTE')
        geolocation_id = tte_geolocation_api_post(ttesession,convention_info)
    return(geolocation_id)

# -----------------------------------------------------------------------
# Create a new location
# -----------------------------------------------------------------------
def tte_geolocation_api_post(ttesession,new_convention):
    geolocation_url = '/api/geolocation'
    geolocation_params = {'session_id': ttesession['id'], 'name': new_convention['location']}
    geolocation_response = requests.post('https://tabletop.events' + geolocation_url, params= geolocation_params)
    geolocation_json = geolocation_response.json()
    geolocation_id = geolocation_json['result']['id']
    return(geolocation_id)

# -----------------------------------------------------------------------
# Get a location name from a tte ID
# -----------------------------------------------------------------------
def tte_geolocation_byid_api_get(ttesession):
    geolocation_url = '/api/geolocation/' + tteconvention_data['result']['geolocation_id']
    geolocation_params = {'session_id': ttesession['id']}
    geolocation_response = requests.get('https://tabletop.events' + geolocation_url, params= geolocation_params)
    geolocation_json = geolocation_response.json()
    geolociation_name = geolocation_json['result']['name']
    return(geolociation_name)

# -----------------------------------------------------------------------
# Routes
# -----------------------------------------------------------------------
# -----------------------------------------------------------------------
# Login to server route
# -----------------------------------------------------------------------
@app.route('/login', methods=['POST', 'GET'])
def login():
    form = LoginForm( )
    session['id'] = 0

    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        if form.validate():
            testuser = User.query.filter_by(username=username).first()
            try:
                if username == testuser.username:
                    if bcrypt.check_password_hash(testuser.password,password):
                        session['name'] = testuser.name
                        session['role'] = testuser.role
                        session['ttesession'] = tte_session()
                        return redirect(url_for('index'))
                    else:
                        flash('Incorrect Password entered')
                        return redirect(request.url)
            except AttributeError:
                flash('Unable to find user')
                return redirect(request.url)
        else:
            flash('All the form fields are required.')
            return redirect(request.url)
    return render_template('login.html', form=form)

# -----------------------------------------------------------------------
# New User Route
# -----------------------------------------------------------------------
@app.route('/newuser', methods=['GET', 'POST'])
def newuser():
    new_user = User()
    createuserform = CreateUserForm(request.form)
    if 'name' in session:
        name = session.get('name')
        role = session.get('role')

        if request.method == 'POST':
            name = request.form['name']
            username = request.form['username']
            password = request.form['password']
            email = request.form['email']
            role = request.form['role']
            if createuserform.validate():
                new_user.name = name
                new_user.username = username
                new_user.password = bcrypt.generate_password_hash(password).decode('utf-8')
                print (new_user.password)
                new_user.email = email
                new_user.role = role
                new_user.id = str(uuid.uuid4())
                try:
                    db.session.add(new_user)
                    db.session.commit()
                    flash('User Saved')
                    return redirect(request.url)
                except:
                    flash('Unable to save user')
                    return redirect(request.url)
    return render_template('newuser.html', createuserform = createuserform, **{'name' : name, 'role' : role})

# -----------------------------------------------------------------------
# Reset Password Route
# -----------------------------------------------------------------------
@app.route('/passwordreset', methods=['GET', 'POST'])
def passwordreset():
    user = User()
    resetpasswordform = ResetUserForm(request.form)

    if 'name' in session:
        name = session.get('name')
        try:
            existing_user = User.query.filter_by(name=name).first()
            print (existing_user.password)
        except:
            flash("Somehow this user doesn't exist")
            return redirect(request.url)

        if request.method == 'POST':
            if resetpasswordform.validate():
                oldpassword = request.form['oldpassword']
                passwordcheck = request.form['passwordcheck']
                newpassword = request.form['newpassword']
                if bcrypt.check_password_hash(existing_user.password,oldpassword) and bcrypt.check_password_hash(existing_user.password,passwordcheck) and bcrypt.check_password_hash(existing_user.password,newpassword) is False:
                    try:
                        existing_user.password = bcrypt.generate_password_hash(newpassword).decode('utf-8')
                        db.session.commit()
                        flash('User Saved')
                        return redirect(request.url)
                    except:
                        flash('Unable to save new password')
                        return redirect(request.url)
                elif bcrypt.check_password_hash(existing_user.password,passwordcheck) is False or bcrypt.check_password_hash(existing_user.password,passwordcheck) is False:
                    flash('Your passwords do not match')
                    return redirect(request.url)
                elif bcrypt.check_password_hash(existing_user.password,newpassword):
                    flash('Your new password matches your old password, please enter in a new password')
                    return redirect(request.url)
                else:
                    flash('Something is not right')
                    return redirect(request.url)
    return render_template('passwordreset.html', resetpasswordform = resetpasswordform, **{'name' : name})

# -----------------------------------------------------------------------
# Upload file Route
# -----------------------------------------------------------------------
@app.route('/upload', methods=['GET', 'POST'])
# Display a visual to upload a CSV file.
def upload():
    folder=config.UPLOAD_FOLDER

    if request.method == 'POST':
        # check if the post request has the file part
        if 'file' not in request.files:
            flash('No file part')
            return redirect(request.url)
        file = request.files['file']
        # if user does not select file, browser also
        # submit an empty part without filename
        if file.filename == '':
            flash('No selected file')
            return redirect(request.url)
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file.save(os.path.join(folder, filename))
            return render_template('upload.html', filename=filename)
    return render_template('upload.html')

# -----------------------------------------------------------------------
# Upload file Route
# -----------------------------------------------------------------------
@app.route('/downloads/<path:filename>', methods=['GET'])
def download(filename):
    return send_from_directory(directory=os.getcwd()+".downloads", filename = filename)

# -----------------------------------------------------------------------
# New Convention Route
# -----------------------------------------------------------------------
@app.route('/newconvention', methods=['GET', 'POST', 'PUT'])
def newconvention():
    # Declarations
    new_convention = {}
    # Call the global so we can modify it in the function with the API call.
    ttesession = session.get('ttesession')
    # Form Function calls
    newconventionform = NewConventionForm(request.form)
    if request.method == "POST":
        print (request.form)
        new_convention['name'] = request.form['name']
        new_convention['location'] = request.form['location']
        new_convention['description'] = request.form['description']
        new_convention['phone_number'] = request.form['phone_number']
        new_convention['email'] = request.form['email']
        new_convention['dates'] = request.form['dates']
        if newconventionform.validate():
            print ('Creating Convention')
            created_convention = tte_convention_convention_api_post(ttesession,new_convention)
            print ('Convention ', new_convention['name'], created_convention, ' created')
            return render_template('newconvention.html', newconventionform=newconventionform)
    return render_template('newconvention.html', newconventionform=newconventionform)

# -----------------------------------------------------------------------
# Conventions Route
# -----------------------------------------------------------------------
@app.route('/conventions', methods=['GET', 'POST', 'PUT'])
def conventions():
    # Declarations
    global tteconvention_data
    name = session.get('name')
    ttesession = session.get('ttesession')
    folder = config.UPLOAD_FOLDER
    files = os.listdir(folder)
    tteconventions = gettteconventions(ttesession)
    # Form Function calls
    conform = ConForm(request.form, obj=tteconventions)
    conform.selectcon.choices = [(tteconventions[con]['id'],tteconventions[con]['name']) for con in tteconventions]
    fileform = FileForm(request.form, obj=files)
    fileform.selectfile.choices = [(file,file) for file in files]
    if request.method == "POST":
        # Pull all the data regarding the convention
        if request.form.get('consubmit'):
            all_days = []
            session['tteconvention_id'] = request.form.get('selectcon',None)
            print ('Getting Convention Information')
            tte_convention_api_get(ttesession,session['tteconvention_id'])
            updateconform = conform_info()
            print (ttesession,session['tteconvention_id'])
            return render_template('conventions.html', updateconform=updateconform, conform=conform, fileform=fileform, **{'name' : name,
            'tteconventions' : tteconventions,
            'tteconvention_data' : tteconvention_data,
            })
        # Submit file to setup the volunteer application
        if request.form.get('conventionsubmit') and session.get('tteconvention_id') is not None:
            update_convention = {}
            print ('Updating the convention')
            update_convention['name'] = request.form['name']
            update_convention['location'] = request.form['location']
            update_convention['description'] = request.form['description']
            update_convention['email'] = request.form['email']
            update_convention['phone_number'] = request.form['phone_number']
            update_convention['dates'] = request.form['dates']
            tte_convention_convention_api_put(ttesession,update_convention)
            updateconform = conform_info()
            return render_template('conventions.html', updateconform=updateconform, conform=conform, fileform=fileform, **{'name' : name,
            'tteconventions' : tteconventions,
            'tteconvention_data' : tteconvention_data,
            })
        # Create a volunter report for event coord to use
        if request.form.get('volunteerreport') and session.get('tteconvention_id') is not None:
            tteconvention_id = session.get('tteconvention_id')
            tteconvention_name = tteconvention_data['result']['name']
            filename = create_volunteer_report(ttesession,session['tteconvention_id'])
            updateconform = conform_info()
            return redirect(url_for('download',filename=filename))
        # Create a CSV file on volunteers for event coord to use
        if request.form.get('volunteercsv') and session.get('tteconvention_id') is not None:
            tteconvention_id = session.get('tteconvention_id')
            tteconvention_name = tteconvention_data['result']['name']
            volunteer_data_csv(tteconvention_data['volunteers'])
            updateconform = conform_info()
            filename = 'volunteerdata.csv'
            return redirect(url_for('download',filename=filename))
        # Create a CSV file on events for event coord to use
        if request.form.get('eventcsv') and session.get('tteconvention_id') is not None:
            tteconvention_id = session.get('tteconvention_id')
            tteconvention_name = tteconvention_data['result']['name']
            event_data_csv(tteconvention_data['events'])
            updateconform = conform_info()
            filename = 'eventdata.csv'
            return redirect(url_for('download',filename=filename))
        # Updates the convention
        if request.form.get('conventionsave') and session.get('tteconvention_id') is not None:
            tteconvention_id = session.get('tteconvention_id')
            tteconvention_name = tteconvention_data['result']['name']
            # Slot Management
            conventionselect = request.form.get('selectfile')
            location = os.path.join(folder,conventionselect)
            print ('Parsing the Convention Matrix File')
            convention_info = convention_parse(location,tteconvention_id,tteconvention_name)
            print ('Creating the Volunteer Shifts')
            tte_convention_volunteer_shift_api_post(ttesession,tteconvention_id,convention_info)
            updateconform = conform_info()
            return render_template('conventions.html', updateconform=updateconform, conform=conform, fileform=fileform, **{'name' : name,
            'tteconventions' : tteconventions,
            'tteconvention_name' : tteconvention_name,
            'tteconvention_data' : tteconvention_data,
            'convention_info' : convention_info,
            })
        # Deprecated....(I think)
        if request.form.get('eventsave') and session.get('tteconvention_id') is not None:
            tteconvention_id = session.get('tteconvention_id')
            tteconvention_name = tteconvention_data['result']['name']
            eventselect = request.form.get('selectfile')
            location = os.path.join(folder,eventselect)
            savedevents = event_parse(location,tteconvention_id,tteconvention_name)
            pushevents = tte_convention_events_api_post(ttesession,tteconvention_id,savedevents)
            updateconform = conform_info()
            return render_template('conventions.html', updateconform=updateconform, conform=conform, fileform=fileform, **{'name' : name,
            'tteconventions' : tteconventions,
            'tteconvention_name' : tteconvention_name,
            'tteconvention_data' : tteconvention_data,
            'savedevents' : savedevents
            })
        # Delete all Events
        if request.form.get('eventsdelete') and session.get('tteconvention_id') is not None:
            tteconvention_id = session.get('tteconvention_id')
            tteconvention_name = tteconvention_data['result']['name']
            tteevents = tte_events_api_get(ttesession,tteconvention_id)
            deleteevents = tte_convention_events_api_delete(ttesession,tteconvention_id,tteevents)
            updateconform = conform_info()
            return render_template('conventions.html', updateconform=updateconform, conform=conform, fileform=fileform, **{'name' : name,
            'tteconventions' : tteconventions,
            'tteconvention_name' : tteconvention_name,
            'tteconvention_data' : tteconvention_data,
            })
        # Delete all volunteer shifts
        if request.form.get('shiftsdelete') and session.get('tteconvention_id') is not None:
            tteconvention_id = session.get('tteconvention_id')
            tteconvention_name = tteconvention_data['result']['name']
            tteshifts = tte_convention_volunteer_shifts_api_get(ttesession,tteconvention_id)
            deleteshifts = tte_convention_volunteer_shifts_api_delete(ttesession,tteconvention_id,tteshifts)
            convention_info = list_convention_info(tteconvention_id)
            #databaseslotdelete = database_slot_delete(tteconvention_id)
            updateconform = conform_info()
            return render_template('conventions.html', updateconform=updateconform, conform=conform, fileform=fileform, **{'name' : name,
            'tteconventions' : tteconventions,
            'tteconvention_name' : tteconvention_name,
            'tteconvention_data' : tteconvention_data,
            })
        # delete all dayparts
        if request.form.get('daypartsdelete') and session.get('tteconvention_id') is not None:
            tteconvention_id = session.get('tteconvention_id')
            tteconvention_name = tteconvention_data['result']['name']
            ttedayparts = tte_convention_dayparts_api_get(ttesession,tteconvention_id)
            deletedayparts = tte_convention_dayparts_api_delete(ttesession,tteconvention_id,ttedayparts)
            updateconform = conform_info()
            return render_template('conventions.html', updateconform=updateconform, conform=conform, fileform=fileform, **{'name' : name,
            'tteconventions' : tteconventions,
            'tteconvention_name' : tteconvention_name,
            'tteconvention_data' : tteconvention_data,
            })
        # Deletes the rooms and spaces (tables)
        if request.form.get('roomsandtablesdelete') and session.get('tteconvention_id') is not None:
            tteconvention_id = session.get('tteconvention_id')
            tteconvention_name = tteconvention_data['result']['name']
            tterooms = tte_convention_rooms_api_get(ttesession,tteconvention_id)
            ttespace = tte_convention_spaces_api_get(ttesession,tteconvention_id)
            tte_convention_roomnsandspaces_api_delete(ttesession,tteconvention_id,tterooms,ttespace)
            updateconform = conform_info()
            return render_template('conventions.html',  updateconform=updateconform, conform=conform, fileform=fileform, **{'name' : name,
            'tteconventions' : tteconventions,
            'tteconvention_name' : tteconvention_name,
            'tteconvention_data' : tteconvention_data,
            })
        # Deletes a volunteer
        if request.form.get('volunteerdelete') and session.get('tteconvention_id') is not None:
            tteconvention_id = session.get('tteconvention_id')
            tteconvention_name = tteconvention_data['result']['name']
            volunteer_id = request.form.get('volunteer_id')
            tte_convention_volunteer_api_delete(ttesession,tteconvention_id,volunteer_id)
            tteconvention_data['volunteers'] = tte_convention_volunteer_api_get(ttesession,tteconvention_id)
            updateconform = conform_info()
            return render_template('conventions.html',  updateconform=updateconform, conform=conform, fileform=fileform, **{'name' : name,
            'tteconventions' : tteconventions,
            'tteconvention_name' : tteconvention_name,
            'tteconvention_data' : tteconvention_data,
            })
        else:
            return render_template('conventions.html',  updateconform=updateconform, conform=conform, fileform=fileform, **{'name' : name })

    else:
        return render_template('conventions.html', conform=conform, fileform=fileform, **{'name' : name,
        })

# -----------------------------------------------------------------------
# Index Route
# -----------------------------------------------------------------------
@app.route('/', methods=['GET', 'POST'])
def index():
    logout = LogoutForm()
    # Check to see if the user already exists.
    # If it does, pass the user's name to the render_template
    if 'name' in session:
        name = session.get('name')
        role = session.get('role')

        ttesession = session.get('ttesession')
        if request.method == 'POST':
            if request.form.get('logoutsubmit'):
                session.pop('name')
                delete_session_params = {'session_id': ttesession['id']}
                delete_session = requests.delete('https://tabletop.events/api/session/' + ttesession['id'], params= delete_session_params)
                session.pop('ttesession')
                return render_template('base.html')
            else:
                pass
        else:
            return render_template('base.html', logout = logout, **{'name' : name, 'role' : role})
    else:
        return render_template('base.html')

# -----------------------------------------------------------------------
# Run Program
# -----------------------------------------------------------------------
if __name__ == '__main__':
    app.run(port=config.PORT, host=config.HOST)
